from docx import Document
from pathlib import Path
import json

root = Path(__file__).resolve().parent.parent
metrics_path = root / 'models' / 'metrics.json'

with open(metrics_path, 'r', encoding='utf-8') as f:
    metrics = json.load(f)

doc = Document()
doc.add_heading('FreightAI Assessment Report', level=1)
doc.add_paragraph('This report summarizes the FreightAI machine learning assessment deliverables, model performance, and generated outputs.')

doc.add_heading('Model Performance', level=2)
summary = doc.add_paragraph()
summary.add_run(f"Best model: {metrics.get('best_model', 'N/A')}\n").bold = True
summary.add_run(f"Validation MAE: {metrics.get('val_mae', 'N/A')}\n")
summary.add_run(f"Validation RMSE: {metrics.get('val_rmse', 'N/A')}\n")
summary.add_run(f"Validation R2: {metrics.get('val_r2', 'N/A')}\n")
summary.add_run(f"Validation MAPE: {metrics.get('val_mape', 'N/A')}")

doc.add_heading('Generated Deliverables', level=2)
doc.add_paragraph('- Trained model artifacts saved in backend/models')
doc.add_paragraph('- Validation predictions saved in backend/validation_predictions.csv')
doc.add_paragraph('- December forecast predictions saved in backend/december_chart_predictions.csv')
doc.add_paragraph('- Charts generated in backend/charts')
doc.add_paragraph('- FastAPI backend and responsive frontend completed')

out_path = root / 'reports' / 'FreightAI_Assessment_Report.docx'
doc.save(out_path)
print(f'Report saved to {out_path}')
